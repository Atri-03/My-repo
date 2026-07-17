"""Node 3: fs_ts_generation.

Takes the structured requirements JSON and generates Functional Spec (FS) and
Technical Spec (TS) documents as .docx files using python-docx.
"""
from __future__ import annotations

import os

from docx import Document

from app.config import settings
from models.schemas import StructuredRequirements
from nodes.state import PipelineState


def _build_fs_document(requirements: StructuredRequirements) -> Document:
    doc = Document()
    doc.add_heading(f"Functional Specification: {requirements.title}", level=1)
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(requirements.summary or "N/A")

    doc.add_heading("Process Flow", level=2)
    for step in requirements.process_flow:
        doc.add_paragraph(
            f"{step.step_number}. {step.description}"
            + (f" (Actor: {step.actor})" if step.actor else ""),
            style="List Number",
        )

    doc.add_heading("Actors", level=2)
    for actor in requirements.actors:
        doc.add_paragraph(f"{actor.name}: {actor.responsibility}", style="List Bullet")

    doc.add_heading("Fields", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for idx, heading in enumerate(["Field", "Description", "Type", "Mandatory"]):
        header_cells[idx].text = heading
    for field in requirements.fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field.name
        row_cells[1].text = field.description
        row_cells[2].text = field.data_type
        row_cells[3].text = "Yes" if field.mandatory else "No"

    doc.add_heading("Validations", level=2)
    for rule in requirements.validations:
        doc.add_paragraph(rule.description, style="List Bullet")

    doc.add_heading("Open Questions", level=2)
    for question in requirements.open_questions:
        doc.add_paragraph(question, style="List Bullet")

    return doc


def _build_ts_document(requirements: StructuredRequirements) -> Document:
    doc = Document()
    doc.add_heading(f"Technical Specification: {requirements.title}", level=1)
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(requirements.summary or "N/A")

    doc.add_heading("Data Elements", level=2)
    for field in requirements.fields:
        doc.add_paragraph(
            f"{field.name} ({field.data_type})"
            + (f" — source: {field.source}" if field.source else ""),
            style="List Bullet",
        )

    doc.add_heading("Validation Logic", level=2)
    for rule in requirements.validations:
        applies = ", ".join(rule.applies_to) if rule.applies_to else "N/A"
        doc.add_paragraph(f"{rule.description} (applies to: {applies})", style="List Bullet")

    doc.add_heading("Proposed Development Object", level=2)
    doc.add_paragraph(
        "A single ABAP report/class skeleton implementing the above process flow "
        "and validations (see generated ABAP source for the POC draft)."
    )

    return doc


def fs_ts_generation(state: PipelineState) -> PipelineState:
    requirements = StructuredRequirements(**state.get("requirements", {}))

    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)
    run_id = state["run_id"]

    fs_path = os.path.join(output_dir, f"{run_id}_FS.docx")
    ts_path = os.path.join(output_dir, f"{run_id}_TS.docx")

    _build_fs_document(requirements).save(fs_path)
    _build_ts_document(requirements).save(ts_path)

    return {
        **state,
        "fs_path": fs_path,
        "ts_path": ts_path,
        "status": "fs_ts_generated",
    }
